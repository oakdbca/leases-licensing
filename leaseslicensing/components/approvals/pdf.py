from io import BytesIO
import os
import re
from django.core.files import File
from django.conf import settings

from docx import Document

from leaseslicensing.components.approvals.models import (
    ApprovalDocument,
)


def docx_replace_regex(
    doc_obj, regex, replace, key, bold=False, italic=False, underline=False, **kwargs
):
    """
    Replaces occurrences of `key` in document paragraphs with value `replace`
    """

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            # Replace the pattern with itself. This causes the entire paragraph to be the run
            # allowing to sub-replace the pattern, but removing any formatting of that paragraph
            # Adding placeholder to table cells in the template preserves the formatting in cells
            p.text = regex.sub(regex.pattern, p.text)
            # TODO Alternatively, can add sophisticated code to split the paragraph into runs
            # first and `add_run` combine the pieces while preserving any formatting and adding
            # inline formatting to the keys. Such code would need to take handle keys that span
            # multiple runs, e.g. `{{key}}` may translate to three runs `{{`, `key`, and `}}`
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                    inline[i].bold = bold
                    inline[i].italic = italic
                    inline[i].underline = underline

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace, key, bold, italic, underline)


def approval_buffer(approval):
    """
    Test function to create a very basic Approval document from template
    """

    licence_doc = f"{settings.BASE_DIR}/leaseslicensing/templates/doc/leases_licence_template.docx"

    doc = Document(licence_doc)

    proposal = approval.current_proposal

    key = "{{applicant_name}}"
    # Licensee example bold (table cell)
    docx_replace_regex(
        doc,
        re.compile(rf"{key}"),
        f"{proposal.applicant_name}",
        key,
        True,
        False,
        False,
    )
    # Licence number example italic  (table cell)
    key = "{{lodgement_number}}"
    docx_replace_regex(
        doc,
        re.compile(rf"{key}"),
        f"{approval.lodgement_number}",
        key,
        False,
        True,
        False,
    )
    # Start/End date examples underlined (no table, inline replacement)
    key = "{{start_date}}"
    start_date = approval.start_date.strftime("%a %d %b %Y, %I:%M %p")
    docx_replace_regex(
        doc, re.compile(rf"{key}"), f"{start_date}", key, False, False, True
    )
    key = "{{expiry_date}}"
    expiry_date = approval.expiry_date.strftime("%a %d %b %Y, %I:%M %p")
    docx_replace_regex(
        doc, re.compile(rf"{key}"), f"{expiry_date}", key, False, False, True
    )
    # Issue date example no formatting (no table, inline replacement)
    key = "{{issue_date}}"
    issue_date = approval.issue_date.strftime("%m/%d/%Y, %H:%M:%S")
    docx_replace_regex(doc, re.compile(rf"{key}"), f"{issue_date}", key, False)

    temp_directory = f"{settings.BASE_DIR}/tmp"

    try:
        os.stat(temp_directory)
    except:
        os.mkdir(temp_directory)
    new_doc_file = f"{temp_directory}/licence_{str(approval.lodgement_number)}.docx"
    new_pdf_file = f"{temp_directory}/licence_{str(approval.lodgement_number)}.pdf"
    doc.save(new_doc_file)
    os.system(
        "libreoffice --headless --convert-to pdf "
        + new_doc_file
        + " --outdir "
        + temp_directory
    )

    approval_buffer = None
    with open(new_pdf_file, "rb") as f:
        approval_buffer = f.read()
    os.remove(new_doc_file)
    os.remove(new_pdf_file)

    approval_buffer = BytesIO(approval_buffer)

    return approval_buffer

def create_approval_doc(approval, **kwargs):
    version_comment = kwargs.get("version_comment", "")
    buffer = approval_buffer(approval)

    filename = "{}.pdf".format(approval.lodgement_number)
    document = ApprovalDocument.objects.create(approval=approval, name=filename)
    document._file.save(filename, File(buffer), save=True)

    buffer.close()

    # Update approval doc filename
    approval.licence_document = document
    approval.save(
        version_comment="Created Approval PDF: {}".format(approval.licence_document.name)
    )

    return document
