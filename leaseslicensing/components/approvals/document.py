from io import BytesIO
import os
import re
import logging
from django.core.files import File
from django.conf import settings
from django.db import IntegrityError, transaction

from docx import Document

from leaseslicensing.components.approvals.models import (
    ApprovalDocument,
)
from leaseslicensing.components.main.decorators import basic_exception_handler

logger = logging.getLogger(__name__)


class ApprovalDocumentGenerator:
    """
    Class to handle approval document generation
    """

    _license_templates = {
        # Test template for fn _example_approval_document_from_template
        "default type": {
            "default_document": f"{settings.BASE_DIR}/leaseslicensing/templates/doc/leases_licence_template.docx"
        },
    }

    def _docx_replace_regex(
        self,
        doc_obj,
        regex,
        replace,
        key,
        bold=False,
        italic=False,
        underline=False,
        **kwargs,
    ):
        """
        Replaces occurrences of `key` in document paragraphs with value `replace`
        """

        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                # Replace the pattern with itself. This causes the entire paragraph to be the run
                # allowing to sub-replace the pattern, but removing any formatting of that paragraph
                # Adding placeholder to table cells in the template preserves the formatting in cells
                p.text = regex.sub(regex.pattern, p.text)
                # TODO Alternatively, can add sophisticated code to split the paragraph into runs
                # first and `add_run` combine the pieces while preserving any formatting and adding
                # inline formatting to the keys. Such code would need to take handle keys that span
                # multiple runs, e.g. `{{key}}` may translate to three runs `{{`, `key`, and `}}`
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
                        inline[i].bold = bold
                        inline[i].italic = italic
                        inline[i].underline = underline

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._docx_replace_regex(
                        cell, regex, replace, key, bold, italic, underline
                    )

    @transaction.atomic
    def _approval_document_for_filename(self, approval, filename_prefix):
        """
        Gets or creates an ApprovalDocument for the given approval
        """

        try:
            document, created = ApprovalDocument.objects.get_or_create(
                approval=approval,
                name__startswith=f"{filename_prefix}{approval.lodgement_number}",
            )
        except IntegrityError as e:
            logger.exception(e)
            raise e
        except Exception as e:
            logger.exception(e)
            raise e
        else:
            return document, created

    @basic_exception_handler
    def _filepath_to_buffer(self, filepath):
        try:
            with open(filepath, "rb") as f:
                buffer = f.read()
        except IOError as e:
            logger.exception(f"Can not open file {filepath}")
            raise e
        else:
            return BytesIO(buffer)

    @basic_exception_handler
    def _example_approval_document_from_template(self, approval, template):
        """
        Test function to create a very basic Approval license document from a template

        Returns:
            BytesIO object
        """

        doc = Document(template)

        proposal = approval.current_proposal

        key = "{{applicant_name}}"
        # Licensee example bold (table cell)
        self._docx_replace_regex(
            doc,
            re.compile(rf"{key}"),
            f"{proposal.applicant_name}",
            key,
            True,
            False,
            False,
        )
        # Licence number example italic  (table cell)
        key = "{{lodgement_number}}"
        self._docx_replace_regex(
            doc,
            re.compile(rf"{key}"),
            f"{approval.lodgement_number}",
            key,
            False,
            True,
            False,
        )
        # Start/End date examples underlined (no table, inline replacement)
        key = "{{start_date}}"
        start_date = approval.start_date.strftime("%a %d %b %Y, %I:%M %p")
        self._docx_replace_regex(
            doc, re.compile(rf"{key}"), f"{start_date}", key, False, False, True
        )
        key = "{{expiry_date}}"
        expiry_date = approval.expiry_date.strftime("%a %d %b %Y, %I:%M %p")
        self._docx_replace_regex(
            doc, re.compile(rf"{key}"), f"{expiry_date}", key, False, False, True
        )
        # Issue date example no formatting (no table, inline replacement)
        key = "{{issue_date}}"
        issue_date = approval.issue_date.strftime("%m/%d/%Y, %H:%M:%S")
        self._docx_replace_regex(
            doc, re.compile(rf"{key}"), f"{issue_date}", key, False
        )

        temp_directory = f"{settings.BASE_DIR}/tmp"

        try:
            os.stat(temp_directory)
        except:
            os.mkdir(temp_directory)
        new_doc_file = f"{temp_directory}/licence_{str(approval.lodgement_number)}.docx"
        new_pdf_file = f"{temp_directory}/licence_{str(approval.lodgement_number)}.pdf"
        doc.save(new_doc_file)
        os.system(
            "libreoffice --headless --convert-to pdf "
            + new_doc_file
            + " --outdir "
            + temp_directory
        )

        buffer = self._filepath_to_buffer(new_pdf_file)

        os.remove(new_doc_file)
        os.remove(new_pdf_file)

        return buffer

    def has_template(self, approval_type_name, document_name=None):
        """
        Returns True if the approval type has a template
        Args:
            approval_type_name:
                Name of the approval type
                E.g. "License (100)"
            document_name (optional):
                Name of the associated document
                E.g. "Cover Letter" or "License (100)" for the actual license document
                Defaults to the approval type name
        """

        if approval_type_name not in self._license_templates.keys():
            return False

        if document_name is None:
            document_name = approval_type_name

        return document_name in self._license_templates[approval_type_name]

    @basic_exception_handler
    def update_approval_document_file(
        self, approval, buffer, filename_prefix, **kwargs
    ):
        """
        Updates the ApprovalDocument file identified by filename with buffer
        Args:
            approval:
                Approval object
            buffer:
                File buffer object
            filename_prefix:
                Prefix to the name of the ApprovalDocument file
            reason:
                Reason for creating the document. Must be one of the choices in ApprovalDocument.REASON_CHOICES
        Returns:
            ApprovalDocument object
        """

        file = File(buffer)
        filename_prefix = filename_prefix or ""
        reason = kwargs.get("reason", "new").lower()

        # Validate reason
        reasons = {r[0]: r[1] for r in ApprovalDocument.REASON_CHOICES}
        if reason not in reasons.keys():
            raise AttributeError(
                f"Invalid reason `{reason}` for approval document not in {reasons.keys()}."
            )

        # Get the current or create a new approval document
        document, created = self._approval_document_for_filename(
            approval, filename_prefix
        )

        if not created:
            # If the document already exists, update the reason unless it is 'new'
            if reason == "new":
                raise AttributeError(
                    f"Reason cannot be 'new'. ApprovalDocument {document.name} already exists."
                )
            document.reason = reason

        # Save the actual file object
        filename = f"{filename_prefix}{approval.lodgement_number}-{approval.lodgement_sequence}.pdf"
        document._file.save(filename, file, save=True)
        # Save the ApprovalDocument object
        document.name = filename
        version_comment = f"{reasons[reason]} Approval document: {filename}"
        logger.info(version_comment)
        document.save(version_comment=version_comment)

        return document

    def create_or_update_approval_document(
        self, approval, filepath=None, filename_prefix=None, **kwargs
    ):
        """
        Returns an ApprovalDocument named filename from the document at filepath,
        either creating a new one or updating an existing one.
        Args:
            approval:
                Approval object
            filepath:
                Path to the approval document file
            filename_prefix:
                Prefix to the name of the approval document file
        """

        buffer = self._filepath_to_buffer(filepath)
        document = self.update_approval_document_file(
            approval, buffer, filename_prefix, **kwargs
        )
        buffer.close()

        return document

    def preview_approval_document(self, approval, filename_prefix=None, **kwargs):
        raise NotImplementedError(
            "Preview of template-generated documents is not implemented yet."
        )

    def create_license_document_from_template(
        self, approval, filename_prefix=None, **kwargs
    ):
        """
        Creates a license document from a template and attaches it to the approval
        Args:
            approval:
                Approval object
            filename_prefix:
                Prefix to the name of the license document file
        """

        # TODO: Replace with actual approval type name derived from approval
        approval_type_name = "default type"
        document_name = "default document"
        if not self.has_template(approval_type_name, document_name):
            raise AttributeError(
                f"Requested document {document_name} Approval type {approval_type_name} does not have a template."
            )

        # TODO: Branching conditional logics here
        if approval_type_name == "default":
            templates = self._license_templates[approval_type_name]
            # TODO: For-loop over templates or similar
            template = templates[document_name]
            buffer = self._example_approval_document_from_template(approval, template)
        else:
            raise NotImplementedError(
                f"Approval document generation from template for approval type {approval_type_name} is not implemented yet."
            )

        document = self.update_approval_document_file(
            approval, buffer, filename_prefix, **kwargs
        )
        buffer.close()
        # Attach the document to the approval
        approval.licence_document = document

        return document
